# src/docx_parser.py
"""Module trích xuất dữ liệu từ file DOCX"""

import re
from typing import Dict, Any, Optional
from docx import Document
from src.utils import clean_text, parse_number


class DocxParser:
    """Class để parse file DOCX phương án sử dụng vốn"""
    
    def __init__(self, file_path: str):
        """
        Khởi tạo parser
        
        Args:
            file_path: Đường dẫn file DOCX
        """
        self.doc = Document(file_path)
        self.text_content = self._extract_all_text()
        self.tables = self._extract_tables()
    
    def _extract_all_text(self) -> str:
        """Trích xuất toàn bộ text từ document"""
        full_text = []
        for para in self.doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
        return "\n".join(full_text)
    
    def _extract_tables(self) -> list:
        """Trích xuất tất cả bảng từ document"""
        tables_data = []
        for table in self.doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            tables_data.append(table_data)
        return tables_data
    
    def _find_value_after_keyword(self, keywords: list, pattern: str = None) -> Optional[str]:
        """
        Tìm giá trị sau từ khóa
        
        Args:
            keywords: Danh sách từ khóa cần tìm
            pattern: Pattern regex để extract
            
        Returns:
            Giá trị tìm được hoặc None
        """
        for keyword in keywords:
            # Tìm trong text
            regex = rf"{keyword}\s*[:：]?\s*(.+?)(?:\n|$)"
            match = re.search(regex, self.text_content, re.IGNORECASE)
            if match:
                value = clean_text(match.group(1))
                if pattern:
                    pattern_match = re.search(pattern, value)
                    if pattern_match:
                        return pattern_match.group(1)
                return value
        return None
    
    def _find_number_after_keyword(self, keywords: list) -> float:
        """Tìm số sau từ khóa"""
        value = self._find_value_after_keyword(keywords)
        if value:
            # Tìm số trong chuỗi
            numbers = re.findall(r'[\d.,]+', value)
            if numbers:
                return parse_number(numbers[0])
        return 0.0
    
    def extract_customer_info(self) -> Dict[str, str]:
        """Trích xuất thông tin khách hàng"""
        return {
            'name': self._find_value_after_keyword([
                'Họ và tên', 'Tên khách hàng', 'Khách hàng', 'Họ tên'
            ]) or '',
            'cccd': self._find_value_after_keyword([
                'CCCD', 'CMND', 'Số CCCD', 'Số CMND', 'Chứng minh nhân dân'
            ]) or '',
            'address': self._find_value_after_keyword([
                'Địa chỉ', 'Địa chỉ thường trú', 'Nơi cư trú', 'Chỗ ở'
            ]) or '',
            'phone': self._find_value_after_keyword([
                'Số điện thoại', 'Điện thoại', 'SĐT', 'Phone'
            ]) or ''
        }
    
    def extract_loan_info(self) -> Dict[str, Any]:
        """Trích xuất thông tin khoản vay"""
        
        # Tìm mục đích vay
        purpose = self._find_value_after_keyword([
            'Mục đích vay', 'Mục đích sử dụng vốn', 'Mục đích',
            'Vay để', 'Sử dụng vốn để'
        ]) or 'Kinh doanh'
        
        # Tìm các số liệu tài chính
        total_need = self._find_number_after_keyword([
            'Tổng nhu cầu vốn', 'Nhu cầu vốn', 'Tổng vốn cần',
            'Tổng mức đầu tư', 'Vốn đầu tư'
        ])
        
        equity = self._find_number_after_keyword([
            'Vốn đối ứng', 'Vốn tự có', 'Nguồn vốn tự có',
            'Vốn chủ sở hữu'
        ])
        
        loan_amount = self._find_number_after_keyword([
            'Số tiền vay', 'Vốn vay', 'Hạn mức vay', 'Dư nợ vay'
        ])
        
        # Nếu không tìm thấy số tiền vay, tính từ tổng nhu cầu - vốn đối ứng
        if loan_amount == 0 and total_need > 0:
            loan_amount = total_need - equity
        
        interest_rate = self._find_number_after_keyword([
            'Lãi suất', 'Lãi suất vay', 'Lãi suất cho vay', 'LS'
        ])
        
        # Mặc định 8.5% nếu không tìm thấy
        if interest_rate == 0:
            interest_rate = 8.5
        
        loan_term = self._find_number_after_keyword([
            'Thời gian vay', 'Thời hạn vay', 'Kỳ hạn', 'Thời hạn'
        ])
        
        # Mặc định 120 tháng nếu không tìm thấy
        if loan_term == 0:
            loan_term = 120
        
        return {
            'purpose': purpose,
            'total_need': total_need,
            'equity': equity,
            'loan_amount': loan_amount,
            'equity_ratio': (equity / total_need * 100) if total_need > 0 else 0,
            'interest_rate': interest_rate,
            'loan_term': int(loan_term),
            'payment_frequency': 'Tháng'
        }
    
    def extract_collateral_info(self) -> Dict[str, Any]:
        """Trích xuất thông tin tài sản bảo đảm"""
        
        asset_type = self._find_value_after_keyword([
            'Loại tài sản', 'Tài sản', 'TSBĐ', 'Loại TSBĐ'
        ]) or 'Bất động sản'
        
        market_value = self._find_number_after_keyword([
            'Giá trị thị trường', 'Giá thị trường', 'Giá trị',
            'Trị giá tài sản'
        ])
        
        asset_address = self._find_value_after_keyword([
            'Địa chỉ tài sản', 'Vị trí', 'Địa điểm', 'Tọa lạc tại'
        ]) or ''
        
        ltv = self._find_number_after_keyword([
            'LTV', 'Tỷ lệ cho vay', 'Tỷ lệ LTV'
        ])
        
        if ltv == 0:
            ltv = 70.0
        
        legal_docs = self._find_value_after_keyword([
            'Giấy tờ pháp lý', 'Pháp lý', 'Giấy tờ', 'Sổ đỏ'
        ]) or 'Sổ đỏ/Giấy chứng nhận quyền sử dụng đất'
        
        return {
            'asset_type': asset_type,
            'market_value': market_value,
            'asset_address': asset_address,
            'ltv': ltv,
            'legal_docs': legal_docs
        }
    
    def extract_financial_info(self) -> Dict[str, float]:
        """Trích xuất thông tin tài chính bổ sung"""
        
        monthly_income = self._find_number_after_keyword([
            'Thu nhập tháng', 'Thu nhập hàng tháng', 'Doanh thu tháng'
        ])
        
        monthly_expense = self._find_number_after_keyword([
            'Chi phí tháng', 'Chi phí hàng tháng', 'Chi phí'
        ])
        
        other_debt = self._find_number_after_keyword([
            'Nợ khác', 'Công nợ khác', 'Nghĩa vụ nợ khác'
        ])
        
        return {
            'monthly_income': monthly_income,
            'monthly_expense': monthly_expense,
            'other_debt': other_debt
        }
    
    def parse_full_document(self) -> Dict[str, Any]:
        """
        Parse toàn bộ document
        
        Returns:
            Dictionary chứa tất cả thông tin đã trích xuất
        """
        return {
            'customer_info': self.extract_customer_info(),
            'loan_info': self.extract_loan_info(),
            'collateral_info': self.extract_collateral_info(),
            'financial_info': self.extract_financial_info(),
            'raw_text': self.text_content
        }
